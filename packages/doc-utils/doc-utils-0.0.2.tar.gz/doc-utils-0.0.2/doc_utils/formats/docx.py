import string
from typing import List

from doc_utils.formats.base import BaseVariable, BaseDocument
from docx import Document


class Variable(BaseVariable):
    def __init__(self):
        self.runs = []
        self.text = ""
        self.paras = []
        self.pos_start = None
        self.pos_end = None
        self.name = None

    def replace(self, value):
        first = self.runs[0]
        last = self.runs[-1]
        middle = self.runs[1:-1]

        if first != last:
            first.text = first.text[:self.pos_start] + str(value)

            text = last.text
            new_text = text[self.pos_end + 1:]
            last.text = new_text

            if len(last.vars) > 1:
                for variable in last.vars:
                    if variable == self:
                        continue
                    if variable.pos_start > self.pos_end:
                        shift_len = len(text) - len(new_text)
                        variable.pos_start -= shift_len
                        if len(variable.runs) == 1:
                            variable.pos_end -= shift_len

        else:
            text = first.text
            new_text = text[:self.pos_start] + str(value) + text[self.pos_end + 1:]
            first.text = new_text

            if len(first.vars) > 1:
                for variable in first.vars:
                    if variable == self:
                        continue
                    if variable.pos_start > self.pos_end:
                        shift_len = len(text) - len(new_text)
                        variable.pos_start -= shift_len
                        if len(variable.runs) == 1:
                            variable.pos_end -= shift_len

        for x in middle:
            x.text = ""


VAR_NAME_CHARS = string.ascii_letters + string.digits + "_-"


def all_paragraphs(doc):
    for para in doc.paragraphs:
        yield para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para

    for section in doc.sections:
        for para in section.header.paragraphs:
            yield para
        for para in section.footer.paragraphs:
            yield para


def find_variables(doc) -> List[Variable]:
    variables_found = []

    variable = None
    name_parsed = False
    for para in all_paragraphs(doc):
        for r in para.runs:
            text = r.text

            for pos, c in enumerate(text):
                if c == "{":
                    variable = Variable()
                    variable.pos_start = pos

                if variable:
                    if r not in variable.runs:
                        variable.runs.append(r)
                        if not hasattr(r, "vars"):
                            r.vars = []
                        r.vars.append(variable)

                    if para not in variable.paras:
                        variable.paras.append(para)
                    variable.text += c

                    if not name_parsed:
                        if c in VAR_NAME_CHARS:
                            if not variable.name:
                                variable.name = ""
                            variable.name += c

                        elif variable.name is not None:
                            name_parsed = True

                if c == "}":
                    if variable:
                        variable.pos_end = pos
                        variables_found.append(variable)

                        variable = None
                        name_parsed = False

    return variables_found


class DocxDocument(BaseDocument):
    def __init__(self, file_or_path):
        super().__init__(file_or_path)
        self.doc = Document(file_or_path)

    def find_variables(self) -> List[BaseVariable]:
        return find_variables(self.doc)

    def save(self, target):
        self.doc.save(target)



