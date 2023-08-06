from docx import Document

from doc_utils.formats.docx import find_variables


def test_empty():
    doc = Document()
    doc.add_paragraph('A plain paragraph having some ')

    vars = find_variables(doc)

    assert len(vars) == 0


def test_simple_paragraph():
    doc = Document()
    p = doc.add_paragraph('A plain {cat_name} having some ')

    vars = find_variables(doc)
    assert len(vars) == 1

    vars[0].replace("Puf")
    assert vars[0].name == "cat_name"

    assert p.text == 'A plain Puf having some '


def test_heading():
    doc = Document()
    p = doc.add_heading('A plain {cat_name} having some ')

    vars = find_variables(doc)
    assert len(vars) == 1

    vars[0].replace("Puf")
    assert vars[0].name == "cat_name"

    assert p.text == 'A plain Puf having some '


def test_header():
    doc = Document()
    section = doc.sections[0]
    p = section.header.paragraphs[0]
    p.text = 'A plain {cat_name} having some '

    vars = find_variables(doc)
    assert len(vars) == 1

    vars[0].replace("Puf")
    assert vars[0].name == "cat_name"

    assert p.text == 'A plain Puf having some '


def test_footer():
    doc = Document()
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.text = 'A plain {cat_name} having some '

    vars = find_variables(doc)
    assert len(vars) == 1

    vars[0].replace("Puf")
    assert vars[0].name == "cat_name"

    assert p.text == 'A plain Puf having some '


def test_table():
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty {some}'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'

    row_cells = table.add_row().cells
    row_cells[0].text = "val {foo}"
    row_cells[1].text = "123"
    row_cells[2].text = "321"

    vars = find_variables(doc)
    assert len(vars) == 2

    vars[0].replace("lala")
    vars[1].replace("bal")
    assert vars[0].name == "some"
    assert vars[1].name == "foo"

    assert table.rows[0].cells[0].text == "Qty lala"
    assert row_cells[0].text == "val bal"


def test_simple_paragraph_exact():
    doc = Document()
    p = doc.add_paragraph('{cat_name}')

    vars = find_variables(doc)
    assert len(vars) == 1

    vars[0].replace("Puf")
    assert vars[0].name == "cat_name"

    assert p.text == 'Puf'


def test_multiple_in_one_paragraph_simple():
    doc = Document()
    p = doc.add_paragraph('{cat_name}{age}')

    vars = find_variables(doc)
    assert len(vars) == 2
    assert vars[0].name == "cat_name"
    assert vars[1].name == "age"

    vars[0].replace("Puf")
    vars[1].replace(10)

    assert p.text == 'Puf10'


def test_multiple_in_one_paragraph():
    doc = Document()
    p = doc.add_paragraph('A plain {cat_name} of {age} having some ')

    vars = find_variables(doc)
    assert len(vars) == 2
    assert vars[0].name == "cat_name"
    assert vars[1].name == "age"

    vars[0].replace("Puf")
    vars[1].replace(10)

    assert p.text == 'A plain Puf of 10 having some '


def test_multiple_runs():
    doc = Document()
    p = doc.add_paragraph('A plain {cat_name} of {age')
    p.add_run(' bold')
    p.add_run(' and some ')
    p.add_run('} italic. lala').italic = True

    vars = find_variables(doc)
    assert len(vars) == 2
    assert vars[0].name == "cat_name"
    assert vars[1].name == "age"

    vars[0].replace("Puf")
    vars[1].replace(10)

    assert len(p.runs) == 4
    assert p.runs[0].text == "A plain Puf of 10"
    assert p.runs[1].text == ""
    assert p.runs[2].text == ""
    assert p.runs[3].text == " italic. lala"
    assert p.text == 'A plain Puf of 10 italic. lala'


def test_multiple_runs_more_vars():
    doc = Document()
    p = doc.add_paragraph('A plain {cat_name} of {age')
    p.add_run(' bold')
    p.add_run(' and some ')
    p.add_run('} italic. {hoho} lala').italic = True

    vars = find_variables(doc)
    assert len(vars) == 3
    assert vars[0].name == "cat_name"
    assert vars[1].name == "age"
    assert vars[2].name == "hoho"

    vars[0].replace("Puf")
    vars[1].replace(10)
    vars[2].replace("nothing")

    assert len(p.runs) == 4
    assert p.runs[0].text == "A plain Puf of 10"
    assert p.runs[1].text == ""
    assert p.runs[2].text == ""
    assert p.runs[3].text == " italic. nothing lala"
    assert p.text == 'A plain Puf of 10 italic. nothing lala'
